# =============================================
# 📝 Bid Proposal Maker (Gemini + Streamlit + OCR)
# =============================================

import os
import fitz  # PyMuPDF for PDFs
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
import streamlit as st
import google.generativeai as genai

# ---------------------------
# STEP 1: Configure Gemini API
# ---------------------------
API_KEY = os.getenv("GOOGLE_API_KEY") or st.text_input("Enter your GOOGLE_API_KEY:", type="password")
if not API_KEY:
    st.warning("Please provide a GOOGLE_API_KEY to continue.")
    st.stop()

genai.configure(api_key=API_KEY)
model = genai.GenerativeModel("gemini-1.5-flash")

# ---------------------------
# STEP 2: Define Proposal Maker prompts
# ---------------------------
PROMPTS = {
    "executive": """Write an executive summary for a proposal response. 
    Highlight our certifications (8(a), HUBZone, WOSB), proven performance on federal construction projects, 
    and how we align with the client's mission-critical needs.""",

    "technical": """Draft the Technical Approach section of a proposal. 
    Cover: understanding of scope, methodology, risk management, quality control, and schedule management. 
    Write in a professional, proposal-ready tone.""",

    "capabilities": """Summarize our Core Capabilities and Differentiators tailored to this solicitation. 
    Reference our expertise in construction, renovation, facility management, energy efficiency, 
    and relevant past performance examples.""",

    "compliance": """Extract all compliance requirements from the solicitation and explain 
    how our company meets or exceeds them. Include UFC, OSHA, EPA, and force protection compliance 
    where relevant.""",

    "value_prop": """Write the Value Proposition section of a proposal. 
    Emphasize how SFMS delivers secure, efficient, and sustainable solutions, 
    while minimizing risk and ensuring mission readiness for the client.""",

    "full_proposal": """Prepare a structured draft proposal with the following sections:
    - Executive Summary
    - Understanding of Scope
    - Technical Approach
    - Core Capabilities & Expertise
    - Compliance
    - Past Performance
    - Value Proposition
    Write in a persuasive, government-proposal style."""
}

# ---------------------------
# STEP 3: Initialize session state
# ---------------------------
if "texts" not in st.session_state:
    st.session_state.texts = {}

if "summaries" not in st.session_state:
    st.session_state.summaries = {}

if "followups" not in st.session_state:
    st.session_state.followups = {}

# ---------------------------
# STEP 4: File Upload
# ---------------------------
st.title("📝 Bid Proposal Maker (from Solicitations)")

uploaded_files = st.file_uploader(
    "Upload Solicitation / RFP / SOW Documents",
    type=["pdf", "docx", "jpg", "jpeg", "png"],
    accept_multiple_files=True
)

def extract_text_from_pdf(file):
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text("text")
    return text

def extract_text_from_docx(file):
    doc = DocxDocument(file)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_text_from_image(file):
    image = Image.open(file)
    text = pytesseract.image_to_string(image)
    return text.strip()

if uploaded_files:
    for f in uploaded_files:
        if f.name not in st.session_state.texts:
            if f.name.endswith(".pdf"):
                st.session_state.texts[f.name] = extract_text_from_pdf(f)
            elif f.name.endswith(".docx"):
                st.session_state.texts[f.name] = extract_text_from_docx(f)
            elif f.name.lower().endswith((".jpg", ".jpeg", ".png")):
                st.session_state.texts[f.name] = extract_text_from_image(f)
            st.success(f"✅ Loaded: {f.name}")

# ---------------------------
# STEP 5: Prompt Selection
# ---------------------------
if st.session_state.texts:
    st.markdown("### ✍️ Choose proposal section(s) to generate")

    options = st.multiselect(
        "Select one or more sections",
        list(PROMPTS.keys()) + ["custom"]
    )

    custom_prompt = ""
    if "custom" in options:
        custom_prompt = st.text_area("✍️ Enter your custom prompt")

    if st.button("Generate Proposal Draft"):
        for fname, text in st.session_state.texts.items():
            st.session_state.summaries[fname] = {}

            for opt in options:
                prompt_to_use = custom_prompt if opt == "custom" else PROMPTS[opt]
                response = model.generate_content(prompt_to_use + "\n\n" + text[:15000])
                st.session_state.summaries[fname][opt] = response.text

        st.success("📄 Proposal draft generated!")

# ---------------------------
# STEP 6: Display + Chat Refinement
# ---------------------------
if st.session_state.summaries:
    for fname, summaries in st.session_state.summaries.items():
        st.header(f"📑 Proposal Draft for {fname}")

        for prompt_name, summary in summaries.items():
            st.subheader(f"🔹 {prompt_name.capitalize()} Section")
            st.write(summary)

            # Chat refinement (chatbot style)
            followup_key = f"{fname}_{prompt_name}_chat"
            if followup_key not in st.session_state:
                st.session_state[followup_key] = [
                    {"role": "assistant", "content": summary}
                ]

            # Display chat history
            for msg in st.session_state[followup_key]:
                with st.chat_message(msg["role"]):
                    st.write(msg["content"])

            # Chat input
            if user_query := st.chat_input(f"Ask a follow-up about {fname} ({prompt_name})"):
                st.session_state[followup_key].append({"role": "user", "content": user_query})
                with st.chat_message("user"):
                    st.write(user_query)

                context = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state[followup_key]])
                response = model.generate_content(context)
                answer = response.text

                st.session_state[followup_key].append({"role": "assistant", "content": answer})
                with st.chat_message("assistant"):
                    st.write(answer)

    # ---------------------------
    # STEP 7: Download Word
    # ---------------------------
    if st.button("⬇️ Download Proposal Draft"):
        doc = DocxDocument()
        doc.add_heading("Bid Proposal Draft", 0)
        for fname, summaries in st.session_state.summaries.items():
            doc.add_heading(fname, level=1)
            for prompt_name, summary in summaries.items():
                doc.add_heading(prompt_name.capitalize(), level=2)
                doc.add_paragraph(summary)

                # Include chat refinements if available
                followup_key = f"{fname}_{prompt_name}_chat"
                if followup_key in st.session_state:
                    doc.add_heading("Chat Refinements", level=3)
                    for msg in st.session_state[followup_key]:
                        doc.add_paragraph(f"{msg['role'].capitalize()}: {msg['content']}")

        doc.save("proposal_draft.docx")
        with open("proposal_draft.docx", "rb") as f:
            st.download_button("Download Word File", f, file_name="proposal_draft.docx")
